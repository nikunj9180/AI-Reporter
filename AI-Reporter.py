#░░░████████╗███████╗░█████╗░██╗░░██╗░░░░░░███╗░░██╗██╗██╗░░██╗░░░
#░░░╚══██╔══╝██╔════╝██╔══██╗██║░░██║░░░░░░████╗░██║██║██║░██╔╝░░░
#░░░░░░██║░░░█████╗░░██║░░╚═╝███████║█████╗██╔██╗██║██║█████═╝░░░░
#░░░░░░██║░░░██╔══╝░░██║░░██╗██╔══██║╚════╝██║╚████║██║██╔═██╗░░░░
#░░░░░░██║░░░███████╗╚█████╔╝██║░░██║░░░░░░██║░╚███║██║██║░╚██╗░░░
#░░░░░░╚═╝░░░╚══════╝░╚════╝░╚═╝░░╚═╝░░░░░░╚═╝░░╚══╝╚═╝╚═╝░░╚═╝░░░


import openai
import docx
from docx.shared import Cm
from docx.shared import Pt

openai.api_key= "   HERE   "   # <------API key

img_path = '/Users/nikunjgupta/AI-Reporter/LetterHeadSample.png' # <------ Letter Head Path


messages = [
    {"role": "system", "content": "You are a Report writer."},
]
message0 = "Write a 200-300 word Brief content in past tence for"
starting="conducted by Shrinik Club, the official Computer Science Club of Gl bajaj institute of technology and mangement on"
summary="write a brief overall summary for same which should start with main content & end with a quote."

topic=input("Enter Event:")
brief=input("Enter 1-2 Line Brief of event:")
date =  input("Enter Date:")
time=input("Enter Time: ")
venue= input("Enter Venue : ")
participants= input("Enter No. of participants : ")
f1=input("Want to add Faculty Coordinator(Y/N):")
if(f1=="Y" or f1=="y"):
    f2=input("Anju Mam?(Y/N)")
    if(f2=="Y" or f2=="y"):
        faculty="Mrs. Anju Khatri" # <--------> Pre-defined Coordinator 1
    else:
        faculty=input("Enter Faculty Name:")

names=["Shubhi Singh (President)","Juhi Pathak (Vice-President)","Ansh Goyal (Events Lead)"]
e1=input("Wnat to add Event Coordinator(Y/N):")   # ^--------> Pre-defined Coordinator 2
if(e1=="Y"or e1=="y"):
    e2=input("Want to more name?(Y/N):")
    if(e2=="Y" or e2=="y"):
        e3=int(input("Enter No. of Names you want to add:"))
        print("FORMAT --->> NAME (DESIGNATION)")
        for i in range(e3):
            e4=input("Enter Member Details:")
            names.append(e4)
            
ipath=input("Paste the path for Invitation msg(Path/N):")
ppath=input("Paste the path for Poster(Path/N):")

gpath=[]
i1=int(input("Enter No. of Glimpses:"))
if(i1!=0):
    for i in range(i1):
        i2=input("Glimpse Path :")
        gpath.append(i2)

#Command for AI to Generate Report Content.
message=message0+topic+starting+date+"at"+venue+"with participation of"+participants+"students."+"In this"+brief

doc = docx.Document()
section = doc.sections[0]
section.page_width = docx.shared.Inches(8.27) 
section.page_height = docx.shared.Inches(11.69) 
sections = doc.sections
margins=0.4   #Page Margin (in inche)
for section in sections:
    section.top_margin = docx.shared.Inches(margins)
    section.bottom_margin = docx.shared.Inches(margins)
    section.left_margin = docx.shared.Inches(margins)
    section.right_margin = docx.shared.Inches(margins)


paragraph = doc.add_paragraph()
run = paragraph.add_run()
run.add_picture(img_path, width=docx.shared.Inches(8.0))
paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
font_size = docx.shared.Pt(11)
style = doc.styles['Normal']
font = style.font
font.size = Pt(11)
paragraph_format = style.paragraph_format
paragraph_format.line_spacing = 1.0 # <-------> Line-Spacing of Document

heading = doc.add_heading(topic.upper(), level=1)
heading.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
heading.style.font.size = Pt(28)
heading.style.font.name = 'Times New Roman'

paragraph = doc.add_paragraph("GREATER NOIDA") # <--------> Location
font = paragraph.style.font
font.name = 'Times New Roman'
font.size = Pt(11)


paragraph = doc.add_paragraph(date)
font = paragraph.style.font
font.name = 'Times New Roman'
font.size = Pt(11)

paragraph = doc.add_paragraph("")

for para in doc.paragraphs:
    para.paragraph_format.space_after = docx.shared.Pt(0)

table = doc.add_table(rows=4, cols=2, style='Table Grid')
table.autofit = False
table.width = Cm(20)
table.columns[0].width = Cm(7)
table.columns[1].width = Cm(13)


#<------ Table content
cell = table.cell(0, 0)
cell.text = "Name of Department" 
cell = table.cell(0, 1)
cell.text = "Computer Science and Engineering"  
cell = table.cell(1, 0)
cell.text = "Date and time"
cell = table.cell(1, 1)
cell.text = date+" , "+time
cell = table.cell(2, 0)
cell.text = "Venue"
cell = table.cell(2, 1)
cell.text = venue
cell = table.cell(3, 0)
cell.text = "Participants"
cell = table.cell(3, 1)
cell.text = participants

paragraph = doc.add_paragraph("")

messages.append(
        {"role": "user", "content": message},
    )
chat = openai.ChatCompletion.create(
        model="gpt-3.5-turbo", messages=messages
    )   
reply = chat.choices[0].message.content
messages.append({"role": "assistant", "content": reply})
paragraph = doc.add_paragraph(reply)
font = paragraph.style.font
font.name = 'Times New Roman'
font.size = Pt(11)

paragraph = doc.add_paragraph("")

if(f1=="Y" or f1=="y"):
    paragraph1 = doc.add_paragraph("Faculty Coordinator:") # <--------> Coordinator1
    runs = paragraph1.runs
    midpoint = len(runs) // 2
    for i in range(midpoint, len(runs)):
        runs[i].bold = True
    font = paragraph1.style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    paragraph = doc.add_paragraph(faculty)
    font = paragraph.style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    paragraph = doc.add_paragraph("")

if(e1=="Y" or e1=="y"):
    paragraph1 = doc.add_paragraph("Event Coordinator:")  # <--------> Coordinator2
    runs = paragraph1.runs
    midpoint = len(runs) // 2
    for i in range(midpoint, len(runs)):
        runs[i].bold = True
    font = paragraph1.style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    for i in names:
        paragraph = doc.add_paragraph(i)
        font = paragraph.style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)

    paragraph = doc.add_paragraph("")


paragraph = doc.add_paragraph("Under the guidance of our Respected HOD Computer Science and Engineering")
font = paragraph.style.font      # <--------> Footer Section detail
font.name = 'Times New Roman'
font.size = Pt(12)


paragraph = doc.add_paragraph("Dr. Sansar Singh Chauhan Sir.")
font = paragraph.style.font      # <--------> Footer Section detail
font.name = 'Times New Roman'
font.size = Pt(12)

for para in doc.paragraphs:
    para.paragraph_format.space_after = docx.shared.Pt(0)

doc.add_page_break() # <--------> invitation page begins

doc.add_paragraph("")
heading = doc.add_heading('INVITATION MESSAGE', level=1)
heading.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
heading.style.font.size = Pt(28)
heading.style.font.name = 'Times New Roman'
doc.add_paragraph("")
if(ipath.upper()!="N"):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(ipath,height=Cm(21))
    paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break() # <--------> poster page begins

doc.add_paragraph("")
heading = doc.add_heading('EVENT POSTER', level=1)
heading.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
heading.style.font.size = Pt(28)
heading.style.font.name = 'Times New Roman'
doc.add_paragraph("")
if(ppath.upper()!="N"):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(ppath,height=Cm(22))
    paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break() # <--------> glimpse page begins
doc.add_paragraph("")
heading = doc.add_heading('GLIMPSES', level=1)
heading.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
heading.style.font.size = Pt(28)
heading.style.font.name = 'Times New Roman'
doc.add_paragraph("")
if(i1!=0):
    if(i1<=3):
        h=24/i1
    elif(i1>3 and i1<=6):
        h=8
    else:
        h=10
    for i in gpath:
        paragraph = doc.add_paragraph()
        paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(i, height=Cm(h))

doc.add_page_break() # <--------> final page begins
doc.add_paragraph("")
heading1 = doc.add_heading('OVERALL SYNOPSIS', level=1)
heading1.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
heading1.style.font.size = Pt(28)
heading1.style.font.name = 'Times New Roman'

doc.add_paragraph("")

paragraph = doc.add_paragraph("The Shrinik Club organised "+topic+"under the aegis of the Department of Computer Science & Engineering.")
font = paragraph.style.font        # ^--------> Details
font.name = 'Times New Roman'
font.size = Pt(12)

doc.add_paragraph("")

messages.append(
        {"role": "user", "content": summary},
    )
chat = openai.ChatCompletion.create(
        model="gpt-3.5-turbo", messages=messages
    )   
s = chat.choices[0].message.content
paragraph = doc.add_paragraph(s)
font = paragraph.style.font
font.name = 'Times New Roman'
font.size = Pt(11)

doc.add_paragraph("")

line = doc.add_paragraph('------------------------------------------------')
line.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
runs = line.runs
midpoint = len(runs) // 2
for i in range(midpoint, len(runs)):
    runs[i].bold = True

doc.add_paragraph("")

paragraph2 = doc.add_paragraph("We extend our most sincere thanks to the Head of the Department, Dr. Sansar Singh Chauhan Sir for his immense support who always stands forward in order to organize promote such events to make the students of the college outshine in all departments.")
font = paragraph2.style.font        # ^--------> Details
font.name = 'Times New Roman'
font.size = Pt(12)

doc.add_paragraph("")

paragraph3 = doc.add_paragraph("We are glad the efforts put in by everyone have barred us the fruits. We extend our most sincere thanks to the enare Shrinik family for their incessant support, guidance for the great success of ‘"+topic+".’")
font = paragraph3.style.font            # ^--------> Details
font.name = 'Times New Roman'
font.size = Pt(12)

doc.add_paragraph("")
doc.add_paragraph("")
doc.add_paragraph("")

heading5 = doc.add_heading('WE APPRECIATE YOUR PRESENCE AND PROMISE', level=2)
heading5.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER  # ^--------> Details
heading5.style.font.size = Pt(14)
heading5.style.font.name = 'Times New Roman'
heading0 = doc.add_heading(' TO DELIVER SUCH AMAZING PROGRAMS AHEAD.', level=2)
heading0.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER  # ^--------> Details
heading0.style.font.size = Pt(14)
heading0.style.font.name = 'Times New Roman'

doc.add_paragraph("")
doc.add_paragraph("")
doc.add_paragraph("")

tq = doc.add_heading('THANK YOU!', level=2)
tq.style.font.size = Pt(14)  # ^--------> Details
tq.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
tq.style.font.name = 'Times New Roman'

final = doc.add_heading('TEAM SHRINIK', level=3) # <--------> Details
final.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
final.style.font.size = Pt(23)
final.style.font.name = 'Times New Roman'

for para in doc.paragraphs:
    para.paragraph_format.space_after = docx.shared.Pt(0)
    para.paragraph_format.space_before = docx.shared.Pt(0)


doc.save(topic+".docx")
print("-------------Report Created Succesfully-------------")